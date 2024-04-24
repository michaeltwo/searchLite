from django.shortcuts import render
from django.http import HttpResponse, HttpResponseBadRequest, HttpResponseNotFound
from django.conf import settings
from .constants import ALLOWED_FILE_TYPES
from .models import CorpusFile
from datetime import datetime, timedelta
from nltk.stem import PorterStemmer
from PIL import Image
from docx import Document
from PyPDF2 import PdfReader
from pymongo import MongoClient
from django.utils import timezone
from django.http import FileResponse
from django.shortcuts import get_object_or_404
from django.http import HttpResponse
from .constants import ALLOWED_FILE_TYPES
from .utils import highlight_query_in_document
from django.shortcuts import get_object_or_404
from django.http import HttpResponse

from .models import CorpusFile
from django.http import StreamingHttpResponse
import mimetypes
import pytesseract
import filetype
import hashlib
import shutil
import fitz
import csv
import os
import re

import nltk
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('wordnet')
nltk.download('stopwords')

stemmer = PorterStemmer()

client = MongoClient('localhost', 27017)
db = client['SearchLite']  # Replace 'your_database_name' with your actual database name
postings_collection = db['Postings']  # Collection to store postings


def homepage(request):
    return render(request, 'index.html')

def process_documents(file_hashes):
    print("Running Background tasks")
    for file_hash in file_hashes:
        # Retrieve the file path from the database using file_hash
        corpus_file = CorpusFile.objects.get(file_hash=file_hash)
        file_path = os.path.join(settings.BASE_DIR, 'corpus', corpus_file.stored_file_name)
        # Extract text from the file
        text = extract_text(file_path)
        # Preprocess the text
        cleaned_text = clean_and_stem(text)
        # Update the postings
        update_postings(str(corpus_file.id), cleaned_text)
        # Update the CorpusFile object to mark processing as completed
        #corpus_file.processed = True
        #corpus_file.save()

def upload(request):
    valid_files = []
    invalid_files = []
    file_hashes = []

    if request.method == 'POST' and request.FILES.getlist('file'):
        files = request.FILES.getlist('file')
        for file in files:
            file_info = filetype.guess(file.read())
            if file_info is not None and file_info.mime in ALLOWED_FILE_TYPES:
                file_hash = generate_file_hash(file)
                if not CorpusFile.objects.filter(file_hash=file_hash).exists():
                    # Split the file name at the last occurrence of '.'
                    parts = file.name.rsplit('.', 1)
                    timestamp = datetime.now().strftime('%Y%m%d%H%M%S')
                    stored_file_name = f"{parts[0]}_{timestamp}.{parts[1]}"
                    file_path = os.path.join(settings.BASE_DIR, 'corpus', stored_file_name)
                    # Save the uploaded file directly to the corpus directory
                    with open(file_path, 'wb+') as destination:
                        shutil.copyfileobj(file, destination)
                    # Add file_hash to the list for processing
                    file_hashes.append(file_hash)
                    uploaded_file = CorpusFile(
                        uploaded_file_name=file.name,
                        stored_file_name=stored_file_name,
                        file_type=file_info.mime,
                        file_size=file.size,
                        file_hash=file_hash
                    )
                    uploaded_file.save()
                    valid_files.append(file.name)
                else:
                    invalid_files.append(f"{file.name} (duplicate file)")
            else:
                invalid_files.append(f"{file.name} (invalid file type)")

        message = ''
        if valid_files:
            message += f"Files uploaded successfully: {', '.join(valid_files)}. "
        if invalid_files:
            message += f"Invalid files: {', '.join(invalid_files)}"

        if file_hashes:
            process_documents(file_hashes)

        return render(request, 'upload.html', {'message': message})

    return render(request, 'upload.html')


def search(request):
    if request.method == 'POST':
        search_query = request.POST.get('search_input', '') 
        cleaned_and_stemmed_query = clean_and_stem(search_query)
        
        # Find documents matching the exact sequence of query terms
        result = {}
        if len(cleaned_and_stemmed_query) == 1:
            single_term = cleaned_and_stemmed_query[0]
            term_postings = postings_collection.find_one({"term": single_term})["positions"]
            result = {doc_id: {single_term: positions} for doc_id, positions in term_postings.items()}
        else:
            first_term = cleaned_and_stemmed_query[0]
            if first_term in postings_collection.distinct("term"):
                first_term_postings = postings_collection.find_one({"term": first_term})["positions"]
                for doc_id, positions in first_term_postings.items():
                    final_positions = {term: [] for term in cleaned_and_stemmed_query}
                    # Iterate through the positions of the first term
                    for pos in positions:
                        term_pos = {first_term: pos}
                        match = True
                        # Check if all other terms occur in sequence after the first term
                        for i, term in enumerate(cleaned_and_stemmed_query[1:], start=1):
                            term_postings = postings_collection.find_one({"term": term})["positions"]
                            if doc_id not in term_postings or pos + i not in term_postings[doc_id]:
                                match = False
                                break
                            term_pos[term] = pos + i
                        if match:
                            # Add the positions to the final result for each term
                            for term, term_position in term_pos.items():
                                final_positions[term].append(term_position)
                    # If all terms are present in sequence, add the document to the result
                    if all(len(final_positions[term]) > 0 for term in cleaned_and_stemmed_query):
                        result[doc_id] = final_positions
        
        # Retrieve the matching documents from your CorpusFile model
        matching_documents = CorpusFile.objects.filter(id__in=result.keys())

                # Calculate file sizes in KB and MB, and format the upload date
        for document in matching_documents:
            file_size_bytes = os.path.getsize(os.path.join(settings.BASE_DIR, 'corpus', document.stored_file_name))
            document.file_size_kb = round(file_size_bytes / 1024, 2)
            document.file_size_mb = round(file_size_bytes / (1024 * 1024), 2)
            document.uploaded_date = timezone.localtime(document.uploaded_at).strftime('%Y-%m-%d %H:%M:%S')


        return render(request, 'results.html', {'matching_documents': matching_documents, 'search_query': search_query})
    return render(request, 'search.html')



def results(request):
    pass

def load_document_image(request, document_id):
    document = get_object_or_404(Document, pk=document_id)
    response = FileResponse(document.image, content_type='image/jpeg')  # Assuming image is stored as a FileField
    return response

def view_document(request, doc_id):
    document = CorpusFile.objects.get(id=doc_id)
    query = request.GET.get('query', '')

    file_path = os.path.join(settings.BASE_DIR, 'corpus', document.stored_file_name)
    file_name, file_extension = os.path.splitext(document.stored_file_name)
    
    if file_extension == '.pdf':
        # Render PDF document using PDF.js or any other PDF viewer
        return render(request, 'doc_viewer.html', {'doc_id': doc_id,'query': query})
    elif file_extension == '.docx':
        # Convert Word document to PDF and render using PDF.js or other viewer
        pdf_path = os.path.join(settings.BASE_DIR, 'corpus', f'{file_name}.pdf')
        doc = Document(file_path)
        doc.save(pdf_path)
        return render(request, 'pdf_viewer.html', {'file_path': pdf_path, 'query': query})
    elif file_extension in ['.jpg', '.jpeg', '.png', '.gif']:
        # Render image directly
        return render(request, 'image_viewer.html', {'file_path': file_path, 'query': query})
    elif file_extension == '.html':
        # Render HTML document in an iframe
        return render(request, 'html_viewer.html', {'file_path': file_path, 'query': query})
    else:
        pass

def load_document(request, doc_id):
    # Get the document object based on the doc_id
    document = get_object_or_404(CorpusFile, id=doc_id)
    file_path = os.path.join(settings.BASE_DIR, 'corpus', document.stored_file_name)

    return FileResponse(open(file_path, 'rb'))

def file_iterator(file_path, chunk_size=8192):
    with open(file_path, 'rb') as f:
        while True:
            chunk = f.read(chunk_size)
            if not chunk:
                break
            yield chunk

def view_pdf_document(request, doc_id):
    document = get_object_or_404(CorpusFile, id=doc_id)
    file_path = os.path.join(settings.BASE_DIR, 'corpus', document.stored_file_name)
        # Check if the file exists
    if not os.path.exists(file_path):
        return HttpResponseNotFound("The requested file was not found.")

    # Check if the file is a PDF (optional, but recommended)
    if not file_path.endswith('.pdf'):
        return HttpResponseBadRequest("The requested file is not a PDF.")
    
    # Send PDF file as response
    response = StreamingHttpResponse(file_iterator(file_path), content_type='application/pdf')
    response['Content-Disposition'] = f'inline; filename="{document.stored_file_name}"'
    return response

def view_document_image(request, doc_id):
    # Retrieve the CorpusFile object by its ID
    document = get_object_or_404(CorpusFile, id=doc_id)
    
    # Assuming the image is stored as a FileField in your CorpusFile model
    # Replace 'image' with the actual name of the FileField in your model
    image = document.image  # Replace 'image' with the actual name of the FileField
    
    # Read the image data
    with image.open() as f:
        image_data = f.read()
    
    # Set the content type of the response
    content_type = 'image/jpeg'  # Adjust the content type based on your image format
    
    # Return the image data as an HTTP response
    return HttpResponse(image_data, content_type=content_type)


def generate_file_hash(file):
    try:
        file.seek(0)  # Move the file pointer to the beginning
        file_content = file.read()
        file_hash = hashlib.sha256(file_content).hexdigest()
        file.seek(0)
        return file_hash
    except Exception as e:
        print(f"Error generating file hash: {e}")
        return None
    
def extract_text(file_path):
    if file_path.endswith('.pdf'):
        return extract_text_from_pdf(file_path)
    elif file_path.endswith(('.doc', '.docx')):
        return extract_text_from_docx(file_path)
    elif file_path.endswith('.csv'):
        return extract_text_from_csv(file_path)
    elif file_path.endswith(('.txt', '.text')):
        return extract_text_from_text(file_path)
    elif file_path.endswith(('.jpg', '.jpeg', '.png', '.bmp')):
        return extract_text_from_image(file_path)
    else:
        return ''

def clean_and_stem(document):
    # Remove special characters and punctuation
    cleaned_doc = re.sub(r'[^a-zA-Z\s]', '', document)
    # Convert to lowercase
    cleaned_doc = cleaned_doc.lower()
    # Tokenize by splitting
    tokens = cleaned_doc.split()
    # Stemming
    stemmed_tokens = [stemmer.stem(token) for token in tokens]
    return stemmed_tokens

def update_postings(doc_id, cleaned_text):
    for position, term in enumerate(cleaned_text):
        postings = postings_collection.find_one({"term": term})
        if postings is None:
            postings = {"term": term, "positions": {doc_id: [position]}}
        else:
            positions = postings.get("positions", {})
            positions[doc_id] = positions.get(doc_id, []) + [position]
            postings["positions"] = positions
        postings_collection.replace_one({"term": term}, postings, upsert=True)


def extract_text_from_pdf(file_path):
    text = ''
    try:
        with fitz.open(file_path) as pdf_file:
            for page_num in range(pdf_file.page_count):
                page = pdf_file[page_num]
                text += page.get_text()
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
    return text

def extract_text_from_docx(file_path):
    document = Document(file_path)
    return '\n'.join([paragraph.text for paragraph in document.paragraphs])

def extract_text_from_csv(file_path):
    text = ''
    with open(file_path, 'r') as file:
        reader = csv.reader(file)
        for row in reader:
            text += ', '.join(row) + '\n'
    return text

def extract_text_from_text(file_path):
    with open(file_path, 'r') as file:
        return file.read()

def extract_text_from_image(file_path):
    text = ''
    try:
        text = pytesseract.image_to_string(Image.open(file_path), lang='eng')
    except Exception as e:
        print(f"Error extracting text from image: {e}")
    return text
